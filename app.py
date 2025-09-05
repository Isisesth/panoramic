import io
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime, date, timedelta
from docx import Document

# =========================
# CONFIG
# =========================
st.set_page_config(page_title="Panorama do Cliente", layout="wide")
st.title("📄 Panorama do Cliente")

# -------------------------
# Utils
# -------------------------
def parse_date(val, fallback=None):
    """Converte o valor em date (aceita str/datetime/NaT). Usa dayfirst=True para dd/mm/aaaa."""
    if pd.isna(val) or val is None or val == "":
        return fallback
    try:
        d = pd.to_datetime(val, errors="coerce", dayfirst=True)
        if pd.isna(d):
            return fallback
        return d.date()
    except Exception:
        return fallback

def fmt_date(d):
    """Formata date no padrão dd/mm/aaaa, tratando None."""
    return d.strftime("%d/%m/%Y") if isinstance(d, (datetime, date)) else "—"

def safe_progress_value(percentual):
    """Converte percentual (0-100+) em valor seguro para st.progress (0-1)."""
    if percentual is None:
        return 0.0
    return max(0.0, min(percentual / 100.0, 1.0))

def df_to_csv_bytes(df: pd.DataFrame, include_index: bool = True) -> bytes:
    """Gera bytes CSV (UTF-8 BOM) para download."""
    return df.to_csv(index=include_index).encode("utf-8-sig")

# -------------------------
# Tabelas de referência
# -------------------------
SOL_PRAZO = {
    "FOIA": 30, "I-130": 30, "COS": 30, "B2-EXT": 30, "NPT": 60, "NVC": 60, "K1": 30,
    "WAIVER": 90, "EB2-NIW": 120, "EB1": 90, "E2": 90, "O1": 90, "EB4": 90,
    "AOS": 60, "ROC": 60, "NATZ": 30, "DA": 90, "PIP": 30, "I-90": 30, "COURT": 60,
    "DOM": 30, "SIJS": 30, "VAWA": 90, "T-VISA": 120, "U-VISA": 90, "I-918B": 30,
    "ASYLUM": 120, "PERM": 30, "EB3": 30
}

CASE_STAGES = [
    # FOIA
    "FOIA - FORMS IN PREPARATION AND WAITING FOR WELCOME CALL",
    "FOIA - INTAKE FORM SENT",
    "FOIA - FORMS WAITING FOR CLIENT SIGNATURES",
    "FOIA - WAITING FOR CLIENT FINGERPRINTS",
    "FOIA - READY TO FILE",
    "FOIA - FILED PENDING RESPONSE",
    "FOIA - RESULTS RECEIVED",
    "FOIA - RESULTS REVIEWED BY LAWYER",
    "FOIA - CONSULTATION SCHEDULED",
    "FOIA - PREPARE TO CREATE A CASE AND CLOSE",
    # CASE (amostra ampliada)
    "CASE - ON BOARDING - CEM WELCOME CALL",
    "CASE - WAITING TIME TO START",
    "CASE - WAITING WELCOME LETTER FROM NVC",
    "CASE - WAITING FOR CLIENT EVIDENCE",
    "CASE - ON CREDIT HOLD - ACCOUNTING",
    "CASE - CLIENT FIRED",
    "CASE - CLIENT HOLD REQUEST / WAITING FOR DATE TO RE-START",
    "CASE - FORMS IN PREPARATION OR EVIDENCE IN TRANSLATION",
    "CASE - FORMS WAITING FOR CLIENT SIGNATURE",
    "CASE - WAITING FOR BUSINESS PLAN",
    "CASE - CASE READY TO DRAFT COVER LETTER",
    "CASE - WAITING FOR PSYCHOLOGICAL EVALUATION",
    "CASE - CIVIL DOCUMENTS SUBMITTED",
    "CASE - WAITING FOR AFFIDAVIT",
    "CASE - FINAL PACKAGE IN PREPARATION",
    "CASE - WAITING FOR CEM REVIEW",
    "CASE - WAITING FOR SUPERVISOR REVIEW",
    "CASE - WAITING FOR ATTORNEY REVIEW",
    "CASE - WAITING FOR SENIOR ATTORNEY REVIEW",
    "CASE - FINAL PACKAGE ATTORNEY REVIEW",
    "CASE - REVISIONS RECOMMENDED BY ATTORNEY",
    "CASE - APPROVED BY ATTORNEY",
    "CASE - WAITING CLIENT REVIEW (FINAL PACKAGE/COVER LETTER)",
    "CASE - WAITING FOR PAYMENT OF IMMIGRATION FEE",
    "CASE - READY TO GO",
    "CASE - FILED PENDING RECEIPT",
    "CASE - E-FILED WAITING FOR RECEIPT",
    "CASE - USCIS PENDING DECISION",
    "CASE - NVC PENDING DECISION",
    "CASE - USCIS PENDING DECISION - SUPPLEMENTING",
    "CASE - WAITING FOR INTERVIEW DATE",
    "CASE - INTERVIEW SCHEDULED FILE IN PREPARATION",
    "CASE - INTERVIEW SCHEDULED FILE READY",
    "CASE - REQUEST INTERVIEW RESCHEDULE",
    "CASE - NOID RECEIVED",
    "CASE - NOID WAITING FOR CLIENT EVIDENCE",
    "CASE - NOID DRAFTING RESPONSE",
    "CASE - RFE RECEIVED",
    "CASE - RFE WAITING FOR CLIENT EVIDENCE",
    "CASE - RFE DRAFTING RESPONSE",
    "CASE - RFE WAITING FOR ATTORNEY REVIEW",
    "CASE - WAITING FOR ROAD MAP",
    "CASE - WAITING FOR PORTABILITY STATUS ADJUSTMENT",
    "CASE - AOS ADMINISTRATIVELY CLOSED",
    "CASE - APPROVED",
    "CASE - APPROVED - FOLLOW UP IMMIGRANT FEE PAYMENT",
    "CASE - APPROVED - FOLLOW UP GC DELIVERY",
    "CASE - APPROVED - WAITING FOR ADJUSTMENT OF STATUS",
    "CASE - WAITING FOR PRIORITY DATE",
    "CASE - WAITING FOR WAIVER APPROVAL",
    "CASE - CLIENT CANCELLATION",
    "CASE - DENIED",
    "CASE - DENIED - RETURN TO INTAKE TO DISCUSS NEXT STEP WITH RENATA OR CLIENT",
    "CASE - ADDITIONAL CASE",
    "CASE - PHYSICAL FILE READY TO BE CLOSED",
    "CASE - FILE READY TO BE CLOSED",
    "CASE - CLOSED",
    # PERM (amostra)
    "PERM - WAITING COMPANY EVIDENCE",
    "PERM - SEARCHING/DEFINING O*NET CODE",
    "PERM - PW DRAFTING",
    "PERM - PW WAITING CLIENT SIGNATURE",
    "PERM - PW WAITING FOR ATTORNEY'S APPROVAL",
    "PERM - PW READY TO GO",
    "PERM - PW PENDING DETERMINATION",
    "PERM - PW RFI RECEIVED",
    "PERM - PW RFI PENDING DECISION",
    "PERM - PW DETERMINATED - WAITING WAGE APPROVAL",
    "PERM - ADVERTISING STARTED - PW DETERMINED",
    "PERM - LABOR CERTIFICATION DRAFTING",
    "PERM - LABOR CERTIFICATION TEAM REVIEW",
    "PERM - LABOR CERTIFICATION WAITING CLIENT SIGNATURE",
    "PERM - LABOR CERTIFICATION WAITING ATTORNEY'S APPROVAL",
    "PERM - LABOR CERTIFICATION PENDING DECISION",
    "PERM - LABOR CERTIFICATION RFI RECEIVED",
    "PERM - LABOR CERTIFICATION AUDIT RECEIVED",
    "PERM - LABOR CERTIFICATION AUDIT PENDING DECISION",
    "PERM - LABOR CERTIFICATION APPROVED",
    "PERM - LABOR CERTIFICATION DENIED",
    "PERM - LABOR CERTIFICATION DENIED - REQUEST FOR RECONSIDERATION DRAFTING",
    # COURT (amostra)
    "COURT - WAITING TIME TO START",
    "COURT - WAITING FOR CLIENT EVIDENCE",
    "COURT - WAITING FOR USCIS RCP",
    "COURT - WAITING FOR USCIS DECISION",
    "COURT - WAITING FOR CEM REVIEW",
    "COURT - WAITING FOR ATTORNEY REVIEW",
    "COURT - FP IN PREPARATION",
    "COURT - 42B FORM IN PREPARATION",
    "COURT - WAITING FOR FEE PAYMENT",
    "COURT - ADM. CLOSED",
    "COURT - MTW PENDING DECISION",
    "COURT - ADDITIONAL MOTION FILED - PENDING DECISION",
    "COURT - FP FILED - PENDING DECISION",
    "COURT - BOND HEARING PREPARATION",
    "COURT - PREPARING CASE FOR HEARING",
    "COURT - CASE READY FOR HEARING",
    "COURT - MTW GRANTED",
    "COURT - FP FILED - GRANTED",
    "COURT - CASE GRANTED - WAITING FOR PAPER ORDER",
    "COURT - FP FILED - DENIED",
    "COURT - RETURN TO INTAKE",
]

# -------------------------
# Session State
# -------------------------
if "courses" not in st.session_state:
    st.session_state.courses = [{"curso": "", "universidade": "", "conclusao": datetime.now().date()}]

if "stages" not in st.session_state:
    st.session_state.stages = []

if "df_master" not in st.session_state:
    st.session_state.df_master = None

# =========================
# Helpers de mapeamento de colunas
# =========================
EXPECTED_FIELDS = {
    "Case": ["Case", "Caso", "Nome do Caso", "Client", "Cliente", "Assunto"],
    "Case Number": ["Case Number", "Número do Caso", "CaseNo", "Case_ID", "ID"],
    "Open Date": ["Open Date", "Data de Abertura", "Opened", "Start Date", "Início"],
    "Closed Date": ["Closed Date", "Data de Fechamento", "Encerrado", "End Date", "Fechado em"],
    "Statute of Limitations Date": ["Statute of Limitations Date", "SOL", "SOL Date", "Prazo SOL", "Limitation Date"],
    "SOL Satisfied?": ["SOL Satisfied?", "SOL Satisfied", "SOL Cumprido?", "SOL OK?"],
    "Practice Area": ["Practice Area", "Área", "Area", "Practice"],
    "Case Stage": ["Case Stage", "Stage", "Status", "Fase", "Etapa"],
    "Your Next Event": ["Your Next Event", "Próximo Evento", "Next Event"],
    "Your Next Task": ["Your Next Task", "Próxima Tarefa", "Next Task"],
    "Last Status Update": ["Last Status Update", "Última Atualização"],
    "Fee Structure": ["Fee Structure", "Modelo de Cobrança"],
    "Flat Fee": ["Flat Fee", "Valor Fixo"],
    "Primary Billing Contact": ["Primary Billing Contact", "Contato de Cobrança"],
    "Description": ["Description", "Descrição", "Notes"],
    "Lead Attorney": ["Lead Attorney", "Advogado Responsável", "Attorney"],
}

def suggest_mapping(df_cols, synonyms):
    norm = {c: c.strip().lower() for c in df_cols if isinstance(c, str)}
    for syn in synonyms:
        syn_l = syn.strip().lower()
        for col, col_l in norm.items():
            if col_l == syn_l:
                return col
    for syn in synonyms:
        syn_l = syn.strip().lower()
        for col, col_l in norm.items():
            if syn_l in col_l:
                return col
    return None

def run_column_mapping_ui(df):
    st.markdown("### 🧭 Mapeamento de colunas (ajuste se necessário)")
    cols = list(df.columns)
    options = ["(não usar)"] + cols
    mapping = {}
    col1, col2 = st.columns(2)
    left_keys = list(EXPECTED_FIELDS.keys())[:len(EXPECTED_FIELDS)//2]
    right_keys = list(EXPECTED_FIELDS.keys())[len(EXPECTED_FIELDS)//2:]
    with col1:
        for field in left_keys:
            sug = suggest_mapping(cols, EXPECTED_FIELDS[field]) or "(não usar)"
            mapping[field] = st.selectbox(f"{field}", options, index=options.index(sug), key=f"map_{field}")
    with col2:
        for field in right_keys:
            sug = suggest_mapping(cols, EXPECTED_FIELDS[field]) or "(não usar)"
            mapping[field] = st.selectbox(f"{field}", options, index=options.index(sug), key=f"map_{field}")

    rename_dict = {v: k for k, v in mapping.items() if v != "(não usar)"}
    df2 = df.rename(columns=rename_dict).copy()

    for c in df2.columns:
        if isinstance(c, str):
            df2[c] = df2[c].apply(lambda x: x.strip() if isinstance(x, str) else x)

    for dc in ["Open Date", "Closed Date", "Statute of Limitations Date"]:
        if dc in df2.columns:
            df2[dc] = pd.to_datetime(df2[dc], errors="coerce", dayfirst=True).dt.date

    st.success("✔️ Mapeamento aplicado. Se algo ficar errado, ajuste os selects acima.")
    return df2

# =========================
# Modo de uso
# =========================
mode = st.radio("Como deseja preencher os dados?", ["A partir de arquivo", "Manual"], horizontal=True)

# =========================
# MODO: ARQUIVO
# =========================
df_master = None
case_data = {}
selected_case = None

if mode == "A partir de arquivo":
    st.subheader("📂 Upload de Arquivo de Casos")
    uploaded_file = st.file_uploader("Envie um arquivo XLS, XLSX ou CSV", type=["xls", "xlsx", "csv"])

    if uploaded_file:
        try:
            if uploaded_file.name.lower().endswith((".xls", ".xlsx")):
                raw_df = pd.read_excel(uploaded_file)
            else:
                raw_df = pd.read_csv(uploaded_file)
            st.success("✅ Arquivo carregado com sucesso.")
            st.dataframe(raw_df, use_container_width=True)

            with st.expander("🧭 Ajustar colunas (se seus nomes forem diferentes)"):
                df_master = run_column_mapping_ui(raw_df)
            if df_master is None:
                df_master = raw_df.copy()

            st.session_state.df_master = df_master
        except Exception as e:
            st.error(f"❌ Erro ao processar o arquivo: {e}")

    if st.session_state.df_master is not None:
        df_master = st.session_state.df_master

    if df_master is not None and "Case Number" in df_master.columns:
        st.subheader("🔎 Selecione um Caso")
        case_numbers = df_master["Case Number"].dropna().astype(str).unique().tolist()
        if case_numbers:
            selected_case = st.selectbox("Número do Caso", case_numbers)
        else:
            st.warning("Nenhum 'Case Number' encontrado no arquivo.")

    if selected_case and df_master is not None:
        row = df_master[df_master["Case Number"].astype(str) == str(selected_case)].iloc[0]
        case_data = row.to_dict()

        # Dados do caso (arquivo)
        nome           = case_data.get("Case", "")
        case_number    = case_data.get("Case Number", "")
        practice_area  = case_data.get("Practice Area", "")
        case_stage     = case_data.get("Case Stage", "")
        open_date      = parse_date(case_data.get("Open Date"))
        closed_date    = parse_date(case_data.get("Closed Date"))
        sol_date       = parse_date(case_data.get("Statute of Limitations Date"))
        sol_satisfied  = case_data.get("SOL Satisfied?", "")
        next_event     = case_data.get("Your Next Event", "")
        next_task      = case_data.get("Your Next Task", "")
        last_update    = case_data.get("Last Status Update", "")
        fee_structure  = case_data.get("Fee Structure", "")
        flat_fee       = case_data.get("Flat Fee", "")
        billing_contact= case_data.get("Primary Billing Contact", "")
        description    = case_data.get("Description", "")
        lead_attorney  = case_data.get("Lead Attorney", "")

        st.info(f"📌 Caso selecionado: **{nome or '—'}**")

        # === PROGRESSO (topo) ===
        st.subheader("📊 Progresso do Caso")
        hoje = datetime.now().date()

        if open_date and sol_date:
            dias_totais     = (sol_date - open_date).days
            dias_decorridos = (hoje - open_date).days
            dias_restantes  = (sol_date - hoje).days
        else:
            dias_totais = dias_decorridos = dias_restantes = 0

        if open_date and sol_date and dias_totais > 0:
            percentual = round((dias_decorridos / dias_totais) * 100, 2)
        else:
            percentual = 100.0 if (sol_date and hoje >= sol_date) else 0.0

        c1, c2, c3 = st.columns(3)
        c1.metric("Dias decorridos", f"{max(dias_decorridos, 0)}")
        c2.metric("Dias restantes até SOL", f"{max(dias_restantes, 0)}")
        c3.metric("Progresso", f"{percentual:.2f}%")
        st.progress(safe_progress_value(percentual))

        if open_date and sol_date:
            if dias_restantes < 0:
                st.error("⚠️ O prazo de SOL já expirou!")
            elif dias_restantes <= 5:
                st.warning("⏱️ Atenção: menos de 5 dias restantes para o SOL.")
            else:
                st.success("✔️ Dentro do prazo do SOL.")
        else:
            st.warning("📅 Datas insuficientes para calcular o progresso (Open Date e/ou SOL ausentes).")

        # Detalhes
        st.subheader("📑 Detalhes do Caso")
        colA, colB = st.columns(2)
        with colA:
            st.write(f"**Case Number:** {case_number or '—'}")
            st.write(f"**Open Date:** {fmt_date(open_date)}")
            st.write(f"**Closed Date:** {fmt_date(closed_date)}")
            st.write(f"**Statute of Limitations Date:** {fmt_date(sol_date)}")
            st.write(f"**SOL Satisfied?:** {sol_satisfied or '—'}")
            st.write(f"**Practice Area:** {practice_area or '—'}")
            st.write(f"**Case Stage:** {case_stage or '—'}")
        with colB:
            st.write(f"**Your Next Event:** {next_event or '—'}")
            st.write(f"**Your Next Task:** {next_task or '—'}")
            st.write(f"**Last Status Update:** {last_update or '—'}")
            st.write(f"**Fee Structure:** {fee_structure or '—'}")
            st.write(f"**Flat Fee:** {flat_fee or '—'}")
            st.write(f"**Primary Billing Contact:** {billing_contact or '—'}")
            st.write(f"**Lead Attorney:** {lead_attorney or '—'}")
            st.write(f"**Description:** {description or '—'}")

        # Exportação
        if st.button("📥 Exportar para Word"):
            doc = Document()
            doc.add_heading("Panorama do Caso", 0)
            doc.add_paragraph(f"Case: {nome or '—'}")
            doc.add_paragraph(f"Case Number: {case_number or '—'}")
            doc.add_paragraph(f"Practice Area: {practice_area or '—'}")
            doc.add_paragraph(f"Case Stage: {case_stage or '—'}")
            doc.add_paragraph(f"Open Date: {fmt_date(open_date)}")
            doc.add_paragraph(f"Closed Date: {fmt_date(closed_date)}")
            doc.add_paragraph(f"Statute of Limitations Date: {fmt_date(sol_date)}")
            doc.add_paragraph(f"SOL Satisfied?: {sol_satisfied or '—'}")
            doc.add_paragraph(f"Your Next Event: {next_event or '—'}")
            doc.add_paragraph(f"Your Next Task: {next_task or '—'}")
            doc.add_paragraph(f"Last Status Update: {last_update or '—'}")
            doc.add_paragraph(f"Fee Structure: {fee_structure or '—'}")
            doc.add_paragraph(f"Flat Fee: {flat_fee or '—'}")
            doc.add_paragraph(f"Primary Billing Contact: {billing_contact or '—'}")
            doc.add_paragraph(f"Lead Attorney: {lead_attorney or '—'}")
            doc.add_paragraph(f"Description: {description or '—'}")
            doc.add_paragraph(f"Dias decorridos: {max(dias_decorridos, 0)}")
            doc.add_paragraph(f"Dias restantes até SOL: {max(dias_restantes, 0)}")
            doc.add_paragraph(f"Progresso: {percentual:.2f}%")
            nome_arquivo = f"panorama_{(case_number or 'caso').replace(' ', '_')}.docx"
            doc.save(nome_arquivo)
            st.success(f"Documento exportado com sucesso: {nome_arquivo}")

# =========================
# MODO: MANUAL
# =========================
if mode == "Manual":
    # PRACTICE AREA antes do nome
    st.subheader("⚙️ Configuração do Caso (Manual)")
    tipo_caso = st.selectbox("🗂️ Practice Area", list(SOL_PRAZO.keys()))
    sol_dias = SOL_PRAZO[tipo_caso]
    st.info(f"🕒 Prazo SOL: {sol_dias} dias")

    data_inicio = st.date_input("📅 Data de início do processo", value=datetime.now().date())
    prazo_final = data_inicio + timedelta(days=sol_dias)
    hoje = datetime.now().date()

    # === ESTÁGIOS + PROGRESSO (topo) ===
    st.subheader("📌 Estágios do Caso")
    if st.button("➕ Adicionar Estágio"):
        start_date = st.session_state.stages[-1]["end_date"] if st.session_state.stages else data_inicio
        st.session_state.stages.append({
            "stage": CASE_STAGES[0],
            "start_date": start_date,
            "end_date": start_date,
            "dias": 0
        })

    for idx, item in enumerate(st.session_state.stages):
        st.session_state.stages[idx]["stage"] = st.selectbox(
            f"Estágio {idx+1}", CASE_STAGES,
            index=CASE_STAGES.index(item["stage"]) if item["stage"] in CASE_STAGES else 0,
            key=f"stage_{idx}"
        )
        st.session_state.stages[idx]["start_date"] = st.date_input(
            "Data inicial", value=item["start_date"], key=f"start_{idx}"
        )
        st.session_state.stages[idx]["end_date"] = st.date_input(
            "Data final", value=item["end_date"], key=f"end_{idx}"
        )
        start = st.session_state.stages[idx]["start_date"]
        end = st.session_state.stages[idx]["end_date"]
        dias = (end - start).days if (isinstance(end, date) and isinstance(start, date) and end >= start) else 0
        st.session_state.stages[idx]["dias"] = dias
        st.text(f"⏳ {dias} dias neste estágio")

    st.subheader("📊 Progresso do Caso")
    dias_totais = (prazo_final - data_inicio).days
    dias_decorridos = (hoje - data_inicio).days
    dias_restantes = (prazo_final - hoje).days
    if dias_totais > 0:
        percentual = round((dias_decorridos / dias_totais) * 100, 2)
    else:
        percentual = 100.0 if hoje >= prazo_final else 0.0

    col1, col2, col3 = st.columns(3)
    col1.metric("Dias decorridos", f"{max(dias_decorridos, 0)}")
    col2.metric("Dias restantes até SOL", f"{max(dias_restantes, 0)}")
    col3.metric("Progresso", f"{percentual:.2f}%")
    st.progress(safe_progress_value(percentual))

    if dias_restantes < 0:
        st.error("⚠️ O prazo de SOL já expirou.")
    elif dias_restantes <= 5:
        st.warning("⏱️ Atenção: menos de 5 dias restantes para o SOL.")
    else:
        st.info("✔️ Dentro do prazo do SOL.")

    # Dados do cliente
    st.subheader("👤 Dados do Cliente")
    nome = st.text_input("Nome completo")

    # Cursos dinâmicos
    st.subheader("🎓 Formação Acadêmica")
    if st.button("➕ Adicionar Curso"):
        st.session_state.courses.append({"curso": "", "universidade": "", "conclusao": datetime.now().date()})

    for idx, curso in enumerate(st.session_state.courses):
        st.markdown(f"**Curso {idx+1}**")
        st.session_state.courses[idx]["curso"] = st.text_input("Nome do curso", value=curso["curso"], key=f"curso_{idx}")
        st.session_state.courses[idx]["universidade"] = st.text_input("Universidade", value=curso["universidade"], key=f"universidade_{idx}")
        st.session_state.courses[idx]["conclusao"] = st.date_input("Data de conclusão", value=curso["conclusao"], key=f"conclusao_{idx}")

    # Export (manual)
    if st.button("📥 Exportar Panorama (Manual) para Word"):
        doc = Document()
        doc.add_heading("Panorama do Cliente", 0)
        doc.add_paragraph(f"Practice Area: {tipo_caso}")
        doc.add_paragraph(f"Nome: {nome or '—'}")
        doc.add_paragraph(f"Início do processo: {fmt_date(data_inicio)}")
        doc.add_paragraph(f"Prazo final (SOL): {fmt_date(prazo_final)}")
        doc.add_paragraph(f"Dias decorridos: {max(dias_decorridos, 0)}")
        doc.add_paragraph(f"Dias restantes: {max(dias_restantes, 0)}")
        doc.add_paragraph(f"Progresso: {percentual:.2f}%")

        doc.add_heading("Estágios do Caso", level=1)
        for s in st.session_state.stages:
            doc.add_paragraph(
                f"{s['stage']} | {fmt_date(s['start_date'])} → {fmt_date(s['end_date'])} | {s['dias']} dias"
            )

        doc.add_heading("Cursos", level=1)
        for c in st.session_state.courses:
            doc.add_paragraph(f"{c['curso']} - {c['universidade']} ({fmt_date(c['conclusao'])})")

        nome_arquivo = f"panorama_{(nome or 'cliente').replace(' ', '_')}.docx"
        doc.save(nome_arquivo)
        st.success(f"Documento exportado: {nome_arquivo}")

# =========================
# VISÃO GERAL (Arquivo): Área × Case Stage — CASOS ATIVOS
# =========================
st.subheader("📈 Área × Case Stage — Casos ativos (excluídos Approved/Denied)")

df_used = st.session_state.df_master if st.session_state.df_master is not None else None

if mode == "A partir de arquivo" and df_used is not None and not df_used.empty:
    # Normaliza nomes/valores
    df_norm = df_used.copy()
    df_norm.columns = [c.strip() if isinstance(c, str) else c for c in df_norm.columns]

    if "Practice Area" not in df_norm.columns or "Case Stage" not in df_norm.columns:
        st.error("Colunas necessárias não encontradas: 'Practice Area' e/ou 'Case Stage'.")
    else:
        df_norm["Practice Area"] = df_norm["Practice Area"].astype(str).str.strip()
        df_norm["Case Stage"]    = df_norm["Case Stage"].astype(str).str.strip()

        # --- Filtro de Área ---
        all_areas = sorted([a for a in df_norm["Practice Area"].dropna().unique()])
        selected_areas = st.multiselect("Filtrar por Practice Area", options=all_areas, default=all_areas)

        # --- Monta vis_df (apenas ATIVOS: exclui Approved/Denied) ---
        rows = []
        for _, r in df_norm.iterrows():
            area  = r.get("Practice Area")
            stage = r.get("Case Stage")
            case  = r.get("Case")
            if pd.isna(area) or pd.isna(stage):
                continue
            area = str(area).strip()
            stage = str(stage).strip()
            if selected_areas and area not in selected_areas:
                continue
            if "APPROVED" in stage.upper() or "DENIED" in stage.upper():
                continue
            rows.append({"Practice Area": area, "Case Stage": stage, "Case": str(case).strip() if pd.notna(case) else "(sem nome)"})

        if not rows:
            st.info("Não há casos ativos com os filtros selecionados.")
        else:
            vis_df = pd.DataFrame(rows)

            # Pivot Área × Stage
            pivot = vis_df.pivot_table(
                index="Practice Area",
                columns="Case Stage",
                values="Case",
                aggfunc="count",
                fill_value=0
            )

            # -----------------------------
            # NOVAS OPÇÕES DE GRÁFICO
            # -----------------------------
            viz_type = st.selectbox(
                "Escolha o tipo de gráfico para visualizar Área × Case Stage",
                [
                    "Heatmap (matplotlib)",
                    "Barras empilhadas (matplotlib)",
                    "Barras agrupadas por Área (matplotlib)",
                    "Treemap (Plotly)",
                    "Sunburst (Plotly)",
                    "Bolhas (Plotly Scatter Área × Stage)"
                ]
            )

            # 1) Heatmap
            if viz_type == "Heatmap (matplotlib)":
                data = pivot.values
                n_areas, n_stages = data.shape
                fig_w = max(10, min(0.45 * max(n_stages, 1) + 4, 28))
                fig_h = max(3,  min(0.50 * max(n_areas, 1) + 2, 20))
                fig, ax = plt.subplots(figsize=(fig_w, fig_h))
                im = ax.imshow(data, aspect="auto")
                ax.set_xticks(range(n_stages))
                ax.set_xticklabels(pivot.columns, rotation=90)
                ax.set_yticks(range(n_areas))
                ax.set_yticklabels(pivot.index)
                ax.set_xlabel("Case Stage")
                ax.set_ylabel("Practice Area")
                ax.set_title("Casos ativos — Heatmap")
                cbar = fig.colorbar(im, ax=ax)
                cbar.set_label("Número de casos", rotation=90)
                # anotações
                for i in range(n_areas):
                    for j in range(n_stages):
                        val = int(data[i, j])
                        if val > 0:
                            ax.text(j, i, str(val), ha="center", va="center", fontsize=8)
                st.pyplot(fig, clear_figure=True)

            # 2) Barras empilhadas por Área (Top N stages)
            elif viz_type == "Barras empilhadas (matplotlib)":
                stage_totals_all = vis_df["Case Stage"].value_counts()
                max_n = max(1, min(20, len(stage_totals_all)))
                top_n = st.slider("Quantos stages exibir (o restante vira 'Outros')",
                                  min_value=1, max_value=max_n, value=min(8, max_n))
                top_stages = stage_totals_all.head(top_n).index.tolist()
                tmp = vis_df.copy()
                tmp["Stage (agrupado)"] = tmp["Case Stage"].apply(lambda s: s if s in top_stages else "Outros")
                stacked = tmp.groupby(["Practice Area", "Stage (agrupado)"])["Case"].count().unstack(fill_value=0)
                # ordena áreas pelo total
                stacked = stacked.loc[stacked.sum(axis=1).sort_values(ascending=False).index]
                fig_h = max(3, 0.55 * len(stacked))
                fig, ax = plt.subplots(figsize=(12, fig_h))
                y = range(len(stacked))
                left = [0] * len(stacked)
                for stg in stacked.columns:
                    vals = stacked[stg].tolist()
                    ax.barh(y, vals, left=left, label=stg)
                    left = [l + v for l, v in zip(left, vals)]
                ax.set_yticks(list(y))
                ax.set_yticklabels(stacked.index)
                ax.invert_yaxis()
                ax.set_xlabel("Número de casos")
                ax.set_title("Casos ativos — Barras empilhadas por Área")
                ax.legend(ncol=2, fontsize=8)
                st.pyplot(fig, clear_figure=True)

            # 3) Barras agrupadas (cada Área = cluster, colunas = Top N stages)
            elif viz_type == "Barras agrupadas por Área (matplotlib)":
                stage_totals_all = vis_df["Case Stage"].value_counts()
                max_n = max(1, min(10, len(stage_totals_all)))  # limitar para legibilidade
                top_n = st.slider("Top N stages (agrupados)", min_value=1, max_value=max_n, value=min(5, max_n))
                top_stages = stage_totals_all.head(top_n).index.tolist()
                tmp = vis_df[vis_df["Case Stage"].isin(top_stages)]
                grp = tmp.groupby(["Practice Area", "Case Stage"])["Case"].count().unstack(fill_value=0)
                areas = list(grp.index)
                x = range(len(areas))
                width = max(0.8 / max(1, len(grp.columns)), 0.1)
                fig, ax = plt.subplots(figsize=(max(10, 1.2 * len(areas)), 6))
                for i, stg in enumerate(grp.columns):
                    ax.bar([p + i * width for p in x], grp[stg].tolist(), width=width, label=stg)
                ax.set_xticks([p + (len(grp.columns) - 1) * width / 2 for p in x])
                ax.set_xticklabels(areas, rotation=45, ha="right")
                ax.set_ylabel("Casos")
                ax.set_title("Casos ativos — Barras agrupadas (Área × Top stages)")
                ax.legend(fontsize=8, ncol=2)
                st.pyplot(fig, clear_figure=True)

            # 4) Treemap (Plotly)
            elif viz_type == "Treemap (Plotly)":
                try:
                    import plotly.express as px
                    treedf = vis_df.groupby(["Practice Area", "Case Stage"])["Case"].count().reset_index(name="count")
                    fig = px.treemap(treedf, path=["Practice Area", "Case Stage"], values="count")
                    fig.update_layout(margin=dict(l=0, r=0, t=30, b=0), title="Casos ativos — Treemap Área → Stage")
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    st.info("Instale o Plotly para ver esta visualização: `py -m pip install plotly`.")

            # 5) Sunburst (Plotly)
            elif viz_type == "Sunburst (Plotly)":
                try:
                    import plotly.express as px
                    sb = vis_df.groupby(["Practice Area", "Case Stage"])["Case"].count().reset_index(name="count")
                    fig = px.sunburst(sb, path=["Practice Area", "Case Stage"], values="count")
                    fig.update_layout(margin=dict(l=0, r=0, t=30, b=0), title="Casos ativos — Sunburst Área → Stage")
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    st.info("Instale o Plotly para ver esta visualização: `py -m pip install plotly`.")

            # 6) Bolhas (Scatter Área × Stage)
            else:
                try:
                    import plotly.express as px
                    bubbles = vis_df.groupby(["Practice Area", "Case Stage"])["Case"].count().reset_index(name="count")
                    fig = px.scatter(bubbles, x="Practice Area", y="Case Stage", size="count", size_max=40)
                    fig.update_layout(margin=dict(l=0, r=0, t=30, b=0), title="Casos ativos — Bolhas (Área × Stage)")
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    st.info("Instale o Plotly para ver esta visualização: `py -m pip install plotly`.")

            # -------------------------
            # Resumos e Downloads
            # -------------------------
            st.markdown("#### 📋 Casos ativos por Área (Approved/Denied excluídos)")
            area_counts = (
                vis_df.groupby("Practice Area")["Case"]
                .count()
                .sort_values(ascending=False)
                .rename("Casos ativos")
                .reset_index()
            )
            st.dataframe(area_counts, use_container_width=True)

            # -------------------------
            # Concentração por Case Stage (ativos) — Total e por Área
            # -------------------------
            st.markdown("#### 🧭 Concentração por Case Stage (ativos) — Total e por Área")

            if vis_df.empty:
                st.info("Sem dados ativos para calcular a concentração por Case Stage.")
            else:
                # Tabela: linhas = Case Stage, colunas = Practice Area (e Total)
                stage_area_pivot = vis_df.pivot_table(
                    index="Case Stage", columns="Practice Area", values="Case",
                    aggfunc="count", fill_value=0
                )

                # Coluna Total (soma das áreas)
                stage_area_pivot["Total"] = stage_area_pivot.sum(axis=1)

                # Coloca 'Total' como primeira coluna
                ordered_cols = ["Total"] + [c for c in stage_area_pivot.columns if c != "Total"]
                stage_area_pivot = stage_area_pivot[ordered_cols]

                # Ordena as linhas por Total (decrescente)
                stage_area_pivot = stage_area_pivot.sort_values(by="Total", ascending=False)

                # Filtro único por Área (opcional)
                areas_cols = [c for c in stage_area_pivot.columns if c != "Total"]
                sel_area = st.selectbox(
                    "Filtrar por Área (opcional)",
                    options=["(Todas)"] + areas_cols
                )

                if sel_area != "(Todas)":
                    tabela_exibir = stage_area_pivot[["Total", sel_area]].copy()
                    tabela_exibir = tabela_exibir.rename(columns={sel_area: f"{sel_area}"})
                else:
                    tabela_exibir = stage_area_pivot.copy()

                # Mostra a tabela (uma linha por Case Stage; colunas = Total e Áreas)
                st.dataframe(tabela_exibir, use_container_width=True)

                # Exemplo: USCIS PENDING DECISION | Total=10 | EB1=4 | EB2=3 | ...

                # Download CSV dessa concentração
                st.download_button(
                    "⬇️ Baixar concentração por Case Stage (CSV)",
                    data=df_to_csv_bytes(tabela_exibir.reset_index(), include_index=False),
                    file_name="concentracao_case_stage_ativos.csv",
                    mime="text/csv"
                )

            # -------------------------
            # Downloads CSV (pivot e resumo por área)
            # -------------------------
            st.markdown("### ⬇️ Downloads")
            c1, c2 = st.columns(2)
            with c1:
                st.download_button(
                    "Baixar pivot Área × Stage (CSV)",
                    data=df_to_csv_bytes(pivot, include_index=True),
                    file_name="pivot_area_x_stage_ativos.csv",
                    mime="text/csv"
                )
            with c2:
                st.download_button(
                    "Baixar resumo por Área (CSV)",
                    data=df_to_csv_bytes(area_counts, include_index=False),
                    file_name="resumo_por_area_ativos.csv",
                    mime="text/csv"
                )
else:
    st.caption("Carregue um arquivo em 'A partir de arquivo' para ver a visão geral Área × Case Stage.")
